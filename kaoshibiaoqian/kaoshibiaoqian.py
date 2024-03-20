# 考室标签
import pandas as pd
from docx import Document

# 读取Excel文件
data = pd.read_excel('数据源.xlsx')  # 假设Excel文件名为'源数据.xlsx'

# 加载Word模板
template_doc = Document('考室标签.docx')  # 假设Word模板文件名为'template.docx'

# 创建一个新文档用于保存最终结果
final_doc = Document()

# 遍历数据中的每条记录
for index, row in data.iterrows():
    # 如果不是第一条记录，添加分页符
    if index > 0:
        final_doc.add_page_break()

    # 复制模板内容到新文档
    for element in template_doc.element.body:
        import copy
        new_element = copy.deepcopy(element)
        final_doc.element.body.append(new_element)

    # 替换模板中的占位符
    for p in final_doc.paragraphs[-len(template_doc.paragraphs):]:  # 只在最后添加的模板单元中替换
        for col in data.columns:
            placeholder = '{' + col + '}'  # 构造占位符
            if placeholder in p.text:
                p.text = p.text.replace(placeholder, str(row[col]))

# 保存最终文档
final_doc.save('final_document.docx')
